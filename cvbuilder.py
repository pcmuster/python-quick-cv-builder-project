from docx import Document
from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

document.add_picture('me.png', width=Inches(1.15)) #add profile picture to the project path and name it as me

#information sector
name = input('Please enter your name: ')
speak('Hello ' + name + ' how are you today?')

speak('Can you please provide me with your phone number')
phone = input('Please enter your phone number: ')
speak('Can you please provide me with your email address, its better to show your email in your CV')
email = input('Please enter your email address: ')

document.add_paragraph(name.upper() + ' | ' + phone + ' | ' + email) #add text to your documents

#about_me

document.add_heading('About me')
speak('tell me about yourself?')
about_me = input('Tell me about yourself? ')
speak('Ok, that look great')
document.add_paragraph(about_me)

 # work experience
document.add_heading('Work Experience ')
p = document.add_paragraph()
speak('Now i need your work Experience')
company = input('Enter company: ')
from_date = input('from date: ')
to_date = input('to date: ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True

experience_detials = input('Decribe your experince at ' + company)
p.add_run(experience_detials)

#more experiences
while True:
   speak('Do you have more Experience you want to enlist?')
   has_more_exerpience = input('Do you have more experiences? Yes or No ' + ' \n ')
   if has_more_exerpience.lower() == 'yes':
      p = document.add_paragraph()
      company = input('Enter company: ')
      from_date = input('from date: ')
      to_date = input('to date: ')
      p.add_run(company + ' ').bold = True
      p.add_run(from_date + '-' + to_date + '\n').italic = True
      experience_detials = input('Decribe your experince at ' + company + ' \n ')
      p.add_run(experience_detials)
   else:
      speak('Ok, its looking good')
      break
   
#adding Skill
document.add_heading('Skills')
speak('Now to enlist your technical skills')
skills = input('Enter your Skills: ')
p =  document.add_paragraph(skills)
p.style = 'List Bullet'
while True:
    more_skill = input('Want to added more Skill? Yes, No ')
    if more_skill.lower() == 'yes':
        skills = input('Enter your Skills: ')
        p =  document.add_paragraph(skills)
        if skills== 'python':
            speak('I see you code in python that great')
        p.style = 'List Bullet'
    else:
        break

speak('Ok your CV is ready now '+name+ ' hope it help you land a job,')
#footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using Akram Elias while trying to learing python"

document.save('cv.docx') # Saving the